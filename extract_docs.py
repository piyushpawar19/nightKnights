import docx
import os

def docx_to_txt(docx_path, txt_path):
    print(f"Converting {docx_path} to {txt_path}...")
    doc = docx.Document(docx_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            fullText.append(" | ".join(row_text))
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(fullText))

base_dir = r"C:\Users\piyush pawar\\.gemini\antigravity\scratch\india_runs_challenge\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge"
out_dir = r"C:\Users\piyush pawar\\.gemini\antigravity\scratch\india_runs_challenge"

docs = ["README.docx", "job_description.docx", "redrob_signals_doc.docx", "submission_spec.docx"]

for doc_name in docs:
    src = os.path.join(base_dir, doc_name)
    dest = os.path.join(out_dir, doc_name.replace(".docx", ".txt"))
    if os.path.exists(src):
        docx_to_txt(src, dest)
    else:
        print(f"File not found: {src}")
